"""Export tailored resumes and cover letters to PDF, DOCX, and Markdown."""
import io
import re
from typing import Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors


def _clean_inline(text: str) -> str:
    """Remove markdown inline syntax."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def markdown_to_pdf(md: str) -> bytes:
    """Convert markdown resume to Harvard-style PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
        rightMargin=0.7*inch, leftMargin=0.7*inch,
        topMargin=0.6*inch, bottomMargin=0.6*inch)

    # Harvard-style styles
    name_style = ParagraphStyle("Name", fontSize=20, fontName="Helvetica-Bold",
        alignment=1, spaceAfter=6, spaceBefore=0)  # centered
    contact_style = ParagraphStyle("Contact", fontSize=9, textColor=colors.HexColor("#444444"),
        alignment=1, spaceAfter=8, spaceBefore=4)  # centered
    section_style = ParagraphStyle("Section", fontSize=10, fontName="Helvetica-Bold",
        spaceBefore=10, spaceAfter=2, textTransform="uppercase")
    job_title_style = ParagraphStyle("JobTitle", fontSize=10, fontName="Helvetica-Bold",
        spaceBefore=6, spaceAfter=0)
    job_meta_style = ParagraphStyle("JobMeta", fontSize=9, textColor=colors.HexColor("#555555"),
        spaceAfter=2, fontName="Helvetica-Oblique")
    bullet_style = ParagraphStyle("Bullet", fontSize=10, leftIndent=12,
        spaceAfter=1, leading=13, bulletIndent=0)
    body_style = ParagraphStyle("Body", fontSize=10, spaceAfter=3, leading=14)
    skills_style = ParagraphStyle("Skills", fontSize=10, spaceAfter=2, leading=13)

    story = []
    lines = md.strip().split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Name (# header)
        if line.startswith("# "):
            name = _clean_inline(line[2:])
            story.append(Paragraph(name, name_style))
            i += 1
            # Next non-empty line = contact info
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and not lines[i].startswith("#"):
                contact = _clean_inline(lines[i])
                story.append(Paragraph(contact, contact_style))
                story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=4))
                i += 1
            continue

        # Section headers (##)
        if line.startswith("## "):
            title = _clean_inline(line[3:]).upper()
            story.append(Paragraph(title, section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceAfter=3))
            i += 1
            continue

        # Bold lines = job titles
        if line.startswith("**") and line.endswith("**"):
            story.append(Paragraph(_clean_inline(line), job_title_style))
            i += 1
            # Check if next line is italic (company/date/location)
            if i < len(lines) and lines[i].strip().startswith("*") and lines[i].strip().endswith("*"):
                story.append(Paragraph(_clean_inline(lines[i]), job_meta_style))
                i += 1
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            text = _clean_inline(line[2:])
            story.append(Paragraph(f"  •  {text}", bullet_style))
            i += 1
            continue

        # Italic lines (standalone meta info)
        if line.startswith("*") and line.endswith("*"):
            story.append(Paragraph(_clean_inline(line), job_meta_style))
            i += 1
            continue

        # Regular text
        story.append(Paragraph(_clean_inline(line), body_style))
        i += 1

    doc.build(story)
    return buffer.getvalue()


def markdown_to_docx(md: str) -> bytes:
    """Convert markdown resume to Harvard-style DOCX."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = md.strip().split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Name
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean_inline(line[2:]))
            run.font.size = Pt(20)
            run.font.bold = True
            i += 1
            # Contact line
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and not lines[i].startswith("#"):
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(_clean_inline(lines[i]))
                cr.font.size = Pt(9)
                i += 1
            continue

        # Section headers
        if line.startswith("## "):
            h = doc.add_heading(_clean_inline(line[3:]).upper(), level=2)
            for run in h.runs:
                run.font.size = Pt(10)
                run.font.bold = True
            i += 1
            continue

        # Bold = job title
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(_clean_inline(line))
            run.font.bold = True
            run.font.size = Pt(10)
            i += 1
            # Italic next = meta
            if i < len(lines) and lines[i].strip().startswith("*") and lines[i].strip().endswith("*"):
                mp = doc.add_paragraph()
                mr = mp.add_run(_clean_inline(lines[i]))
                mr.font.italic = True
                mr.font.size = Pt(9)
                mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                i += 1
            continue

        # Bullets
        if line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(_clean_inline(line[2:]), style="List Bullet")
            i += 1
            continue

        # Italic standalone
        if line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph()
            r = p.add_run(_clean_inline(line))
            r.italic = True
            r.font.size = Pt(9)
            i += 1
            continue

        doc.add_paragraph(_clean_inline(line))
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def resume_to_pdf(resume_data: dict[str, Any], markdown_content: str = "") -> bytes:
    """Generate PDF from structured resume data."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    name_style = ParagraphStyle("Name", parent=styles["Normal"], fontSize=20, fontName="Helvetica-Bold", spaceAfter=4)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9, textColor=colors.gray, spaceAfter=8)
    section_style = ParagraphStyle("Section", parent=styles["Normal"], fontSize=11, fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=2)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=2)
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10, leftIndent=12, spaceAfter=1)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=10, fontName="Helvetica-Bold", spaceAfter=1)
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9, textColor=colors.gray, spaceAfter=2)

    c = resume_data.get("contact", {})

    # Name
    if c.get("name"):
        story.append(Paragraph(c["name"], name_style))

    # Contact line
    contact_parts = [p for p in [c.get("email"), c.get("phone"), c.get("location"), c.get("linkedin")] if p]
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.black))

    # Summary
    if resume_data.get("summary"):
        story.append(Paragraph("SUMMARY", section_style))
        story.append(Paragraph(resume_data["summary"], body_style))

    # Experience
    if resume_data.get("experience"):
        story.append(Paragraph("EXPERIENCE", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        for exp in resume_data["experience"]:
            end = "Present" if exp.get("current") else exp.get("end_date", "")
            story.append(Paragraph(f"{exp.get('title', '')} — {exp.get('company', '')}", subtitle_style))
            story.append(Paragraph(f"{exp.get('start_date', '')} – {end} | {exp.get('location', '')}", meta_style))
            for bullet in exp.get("bullets", []):
                story.append(Paragraph(f"• {bullet}", bullet_style))
            story.append(Spacer(1, 4))

    # Education
    if resume_data.get("education"):
        story.append(Paragraph("EDUCATION", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        for edu in resume_data["education"]:
            story.append(Paragraph(f"{edu.get('degree', '')} in {edu.get('field', '')} — {edu.get('institution', '')}", subtitle_style))
            story.append(Paragraph(edu.get("graduation_date", ""), meta_style))

    # Skills
    skills = resume_data.get("skills", {})
    all_skills = []
    for v in skills.values():
        if isinstance(v, list):
            all_skills.extend(v)
    if all_skills:
        story.append(Paragraph("SKILLS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        story.append(Paragraph(", ".join(all_skills), body_style))

    # Certifications
    if resume_data.get("certifications"):
        story.append(Paragraph("CERTIFICATIONS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        for cert in resume_data["certifications"]:
            story.append(Paragraph(f"• {cert.get('name', '')} — {cert.get('issuer', '')} ({cert.get('date', '')})", bullet_style))

    # Projects
    if resume_data.get("projects"):
        story.append(Paragraph("PROJECTS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        for proj in resume_data["projects"]:
            story.append(Paragraph(proj.get("name", ""), subtitle_style))
            if proj.get("description"):
                story.append(Paragraph(proj["description"], body_style))
            for bullet in proj.get("bullets", []):
                story.append(Paragraph(f"• {bullet}", bullet_style))

    doc.build(story)
    return buffer.getvalue()


def cover_letter_to_pdf(content: str, candidate_name: str = "") -> bytes:
    """Generate PDF from cover letter text."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("CL", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=12)
    story = []
    for para in content.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip().replace("\n", " "), body_style))
    doc.build(story)
    return buffer.getvalue()


def resume_to_docx(resume_data: dict[str, Any]) -> bytes:
    """Generate DOCX from structured resume data."""
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    c = resume_data.get("contact", {})

    # Name
    if c.get("name"):
        h = doc.add_heading(c["name"], 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.size = Pt(20)

    # Contact
    contact_parts = [p for p in [c.get("email"), c.get("phone"), c.get("location"), c.get("linkedin")] if p]
    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def add_section_header(title: str):
        h = doc.add_heading(title, level=2)
        for run in h.runs:
            run.font.size = Pt(11)

    if resume_data.get("summary"):
        add_section_header("Summary")
        doc.add_paragraph(resume_data["summary"])

    if resume_data.get("experience"):
        add_section_header("Experience")
        for exp in resume_data["experience"]:
            end = "Present" if exp.get("current") else exp.get("end_date", "")
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}").bold = True
            doc.add_paragraph(f"{exp.get('start_date', '')} – {end} | {exp.get('location', '')}").runs[0].font.size = Pt(9)
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    if resume_data.get("education"):
        add_section_header("Education")
        for edu in resume_data["education"]:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')} — {edu.get('institution', '')}").bold = True
            doc.add_paragraph(edu.get("graduation_date", "")).runs[0].font.size = Pt(9)

    skills = resume_data.get("skills", {})
    all_skills = []
    for v in skills.values():
        if isinstance(v, list):
            all_skills.extend(v)
    if all_skills:
        add_section_header("Skills")
        doc.add_paragraph(", ".join(all_skills))

    if resume_data.get("certifications"):
        add_section_header("Certifications")
        for cert in resume_data["certifications"]:
            doc.add_paragraph(f"{cert.get('name', '')} — {cert.get('issuer', '')} ({cert.get('date', '')})", style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def cover_letter_to_docx(content: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    for para in content.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
